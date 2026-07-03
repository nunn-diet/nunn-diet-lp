from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ページ設定 A4
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# カラー定数
NAVY   = RGBColor(0x1A, 0x2B, 0x4A)
GOLD   = RGBColor(0xB8, 0x86, 0x0B)
PINK   = RGBColor(0xE8, 0x6A, 0x7A)
LIGHT_PINK = RGBColor(0xFF, 0xF0, 0xF2)
LIGHT_GOLD = RGBColor(0xFF, 0xF8, 0xE7)
LIGHT_NAVY = RGBColor(0xF0, 0xF2, 0xF8)
GRAY   = RGBColor(0x66, 0x66, 0x66)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LINE_COLOR = RGBColor(0xD4, 0xAF, 0x37)

def set_cell_bg(cell, r, g, b):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{r:02X}{g:02X}{b:02X}')
    tcPr.append(shd)

def add_bottom_border(para, color_hex='D4AF37', size=12):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_para_bg(para, r, g, b):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{r:02X}{g:02X}{b:02X}')
    pPr.append(shd)

def heading_para(text, level=1, color=NAVY, bg=None, align=WD_ALIGN_PARAGRAPH.LEFT, size=None):
    p = doc.add_paragraph()
    p.alignment = align
    if bg:
        set_para_bg(p, bg[0], bg[1], bg[2])
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 0:
        run.font.size = Pt(size or 22)
        run.font.name = 'Yu Gothic'
    elif level == 1:
        run.font.size = Pt(size or 16)
        run.font.name = 'Yu Gothic'
    elif level == 2:
        run.font.size = Pt(size or 13)
        run.font.name = 'Yu Gothic'
    else:
        run.font.size = Pt(size or 11)
        run.font.name = 'Yu Gothic'
    return p

def body_para(text, color=None, size=10.5, italic=False, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Yu Mincho'
    run.italic = italic
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def input_line(label=None, lines=2):
    """記入欄（罫線付き）"""
    if label:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(label)
        run.font.size = Pt(10)
        run.font.name = 'Yu Mincho'
        run.font.color.rgb = GRAY
    for _ in range(lines):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.space_after = Pt(2)
        add_bottom_border(p, 'CCCCCC', 6)
        run = p.add_run(' ')
        run.font.size = Pt(18)
    return p

def section_box(title, color_rgb=(0x1A, 0x2B, 0x4A)):
    """セクションタイトルボックス"""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, *color_rgb)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(6)
    cp.paragraph_format.space_after = Pt(6)
    run = cp.add_run(title)
    run.bold = True
    run.font.color.rgb = WHITE
    run.font.size = Pt(13)
    run.font.name = 'Yu Gothic'
    doc.add_paragraph()

def question_block(q_num, question, hint=None, lines=2):
    """質問ブロック"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run_q = p.add_run(f'Q{q_num}. ')
    run_q.bold = True
    run_q.font.size = Pt(10.5)
    run_q.font.color.rgb = NAVY
    run_q.font.name = 'Yu Gothic'
    run_text = p.add_run(question)
    run_text.font.size = Pt(10.5)
    run_text.font.name = 'Yu Mincho'
    if hint:
        run_h = p.add_run(f'\n　{hint}')
        run_h.font.size = Pt(9)
        run_h.font.color.rgb = GRAY
        run_h.font.name = 'Yu Mincho'
        run_h.italic = True
    input_line(lines=lines)

def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(2)

# ===================== 表紙 =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
run = p.add_run('選ばれる人だけが知っている！')
run.font.size = Pt(11)
run.font.color.rgb = GOLD
run.font.name = 'Yu Mincho'
run.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(10)
run = p.add_run('「この人じゃないとダメ」と')
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = NAVY
run.font.name = 'Yu Gothic'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('思われる')
run.font.size = Pt(20)
run.bold = True
run.font.color.rgb = NAVY
run.font.name = 'Yu Gothic'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ポジションの作り方と')
run.font.size = Pt(18)
run.bold = True
run.font.color.rgb = NAVY
run.font.name = 'Yu Gothic'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('強み発見シート')
run.font.size = Pt(24)
run.bold = True
run.font.color.rgb = GOLD
run.font.name = 'Yu Gothic'

spacer(2)

# 罫線
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
add_bottom_border(p, 'D4AF37', 12)

spacer()

# 解決する悩み
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('── こんな悩みを解決します ──')
run.font.size = Pt(11)
run.font.color.rgb = GOLD
run.font.name = 'Yu Mincho'
run.bold = True

spacer()

concerns = [
    '自分の強みがわからない',
    '他の人と差別化できない',
    'お客様に選ばれる理由がほしい',
    '理想のお客様に深く刺さりたい',
]
t = doc.add_table(rows=1, cols=4)
t.style = 'Table Grid'
for i, text in enumerate(concerns):
    cell = t.rows[0].cells[i]
    set_cell_bg(cell, 0xFF, 0xF0, 0xF2)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after = Pt(8)
    run = cp.add_run(text)
    run.font.size = Pt(9.5)
    run.font.name = 'Yu Mincho'
    run.font.color.rgb = NAVY

spacer(2)

# Pointボックス
t2 = doc.add_table(rows=1, cols=1)
t2.style = 'Table Grid'
cell2 = t2.rows[0].cells[0]
set_cell_bg(cell2, 0x1A, 0x2B, 0x4A)
cp2 = cell2.paragraphs[0]
cp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp2.paragraph_format.space_before = Pt(10)
cp2.paragraph_format.space_after = Pt(10)
run1 = cp2.add_run('Point！　')
run1.bold = True
run1.font.color.rgb = GOLD
run1.font.size = Pt(11)
run1.font.name = 'Yu Gothic'
run2 = cp2.add_run('自分の強みを知り、価値を言語化することで\n「あなたにお願いしたい！」と言われる存在になれます')
run2.font.color.rgb = WHITE
run2.font.size = Pt(10.5)
run2.font.name = 'Yu Mincho'

spacer(3)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('このシートの使い方')
run.bold = True
run.font.size = Pt(11)
run.font.color.rgb = NAVY
run.font.name = 'Yu Gothic'

body_para('PART 1〜3を順番に進めてください。「これくらい大したことじゃない」と思う答えほど、実はあなただけの強みです。\n正直に、思いつくままに書き出すことが大切です。', color=GRAY, size=10)

# 改ページ
doc.add_page_break()

# ===================== PART 1 =====================
p_part = doc.add_paragraph()
p_part.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_part.paragraph_format.space_before = Pt(10)
set_para_bg(p_part, 0x1A, 0x2B, 0x4A)
run = p_part.add_run('PART 1　あなたの「強み」を発見する')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = WHITE
run.font.name = 'Yu Gothic'

spacer()

# STEP 1
section_box('STEP 1　「得意」を深掘りする質問', (0xB8, 0x86, 0x0B))

p_note = doc.add_paragraph()
p_note.paragraph_format.left_indent = Cm(0.3)
run = p_note.add_run('「大したことじゃない」と思うことほど、実は強みです。素直に答えてみてください。')
run.font.size = Pt(9.5)
run.font.color.rgb = GRAY
run.font.name = 'Yu Mincho'
run.italic = True

spacer()

questions_step1 = [
    ('何となく人よりうまくできてしまうことは何ですか？', None),
    ('自分では当たり前だと思っているけど、人から驚かれることは何ですか？', None),
    ('人のどんな状態を見ると「放っておけない」と感じますか？', None),
    ('無意識にイラッとしてしまう人の特徴は何ですか？', '逆に自分の価値観のヒントになります'),
    ('人に何かを伝えるとき、つい熱が入ってしまうテーマは何ですか？', None),
    ('「なんでみんなそれで悩むの？」と感じたことは何ですか？', None),
    ('自分が「これだけは譲れない」と思う価値観は何ですか？', None),
    ('周りの人に「こうなってほしい」と思う理想は何ですか？', None),
    ('「これをやっている自分が一番好き」と思える瞬間はどんな時ですか？', None),
    ('あなたの存在が、誰のどんな未来を変えられると思いますか？', None),
]

for i, (q, hint) in enumerate(questions_step1, 1):
    question_block(i, q, hint, lines=2)

spacer()

# STEP 2
section_box('STEP 2　「経験・ストーリー」から強みを見つける', (0xB8, 0x86, 0x0B))
spacer()

questions_step2 = [
    ('過去の自分が一番苦しかった時、何に悩んでいましたか？', None),
    ('そのとき本当は、誰に何を言ってほしかったですか？', None),
    ('今の自分が、当時の自分にかけてあげられる言葉は何ですか？', None),
    ('これまでの人生で乗り越えてきた経験の中で、「これが今の自分を作った」と思えるものは何ですか？', None),
    ('その経験から得た「自分なりの答え」は何ですか？', None),
]

for i, (q, hint) in enumerate(questions_step2, 11):
    question_block(i, q, hint, lines=2)

spacer()

# STEP 3
section_box('STEP 3　「他人の目」から見た自分', (0xB8, 0x86, 0x0B))
spacer()

items_step3 = [
    ('周りからよく言われること（5つ以上書き出す）', 3),
    ('大切にしている価値観（5つ以上）', 3),
    ('無意識にできていること・強み（5つ以上）', 3),
]
for label, lines in items_step3:
    body_para(f'▼ {label}', bold=True, color=NAVY, size=10.5)
    input_line(lines=lines)
    spacer()

# 改ページ
doc.add_page_break()

# ===================== PART 2 =====================
p_part2 = doc.add_paragraph()
p_part2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_part2.paragraph_format.space_before = Pt(10)
set_para_bg(p_part2, 0x1A, 0x2B, 0x4A)
run = p_part2.add_run('PART 2　差別化ポイントを見つける')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = WHITE
run.font.name = 'Yu Gothic'

spacer()

section_box('「自分だけの差別化軸」発見ワーク', (0xE8, 0x6A, 0x7A))

p_note2 = doc.add_paragraph()
p_note2.paragraph_format.left_indent = Cm(0.3)
run = p_note2.add_run('以下の項目を正直に書き出してください。ネガティブな部分こそ、唯一無二のポジションになります。')
run.font.size = Pt(9.5)
run.font.color.rgb = GRAY
run.font.name = 'Yu Mincho'
run.italic = True

spacer()

diff_items = [
    ('▼ 環境（マイナスでもOK）', '例：離婚経験 / 田舎暮らし / ブラック企業出身 / 子育てしながら起業'),
    ('▼ 外見・体型（自信があってもなくてもOK）', '例：産後20kg太った / 40代でも痩せられた / 元ぽっちゃり'),
    ('▼ 人間関係・家族', '例：離婚危機から夫婦関係を修復 / 夫が太っている / ワンオペ育児'),
]
for label, example in diff_items:
    body_para(label, bold=True, color=NAVY, size=10.5)
    p_ex = doc.add_paragraph()
    p_ex.paragraph_format.left_indent = Cm(0.5)
    run_ex = p_ex.add_run(example)
    run_ex.font.size = Pt(9)
    run_ex.font.color.rgb = GRAY
    run_ex.font.name = 'Yu Mincho'
    run_ex.italic = True
    input_line(lines=2)
    spacer()

spacer()

section_box('差別化コンセプトをつくる', (0xE8, 0x6A, 0x7A))
spacer()

body_para('競合が「　　　　　　　　」をやっている中で、', size=11)
body_para('私は「　　　　　　　　　　　　　　　　」という切り口で発信する。', size=11)
spacer()
body_para('私のアカウントのコンセプト（1文で書く）：', bold=True, color=NAVY, size=10.5)
input_line(lines=2)

spacer(2)

# 改ページ
doc.add_page_break()

# ===================== PART 3 =====================
p_part3 = doc.add_paragraph()
p_part3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_part3.paragraph_format.space_before = Pt(10)
set_para_bg(p_part3, 0x1A, 0x2B, 0x4A)
run = p_part3.add_run('PART 3　ポジションを確立する')
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = WHITE
run.font.name = 'Yu Gothic'

spacer()

section_box('ターゲット設定　具体的な1人を書く', (0x1A, 0x2B, 0x4A))
spacer()

target_fields = [
    ('性別・年齢・居住地', 1),
    ('職業・年収・家族構成', 1),
    ('日常の生活・ルーティン', 2),
    ('趣味・興味関心', 1),
    ('抱えている悩み（具体的に）', 3),
]
for label, lines in target_fields:
    body_para(f'▼ {label}', bold=True, color=NAVY, size=10.5)
    input_line(lines=lines)
    spacer()

spacer()

section_box('キャッチコピーをつくる', (0x1A, 0x2B, 0x4A))
spacer()

p_note3 = doc.add_paragraph()
p_note3.paragraph_format.left_indent = Cm(0.3)
run = p_note3.add_run('以下の要素を組み合わせて、あなたのプロフィール名（キャッチコピー）を作りましょう。')
run.font.size = Pt(9.5)
run.font.color.rgb = GRAY
run.font.name = 'Yu Mincho'
run.italic = True

spacer()

catch_fields = [
    '自分の個性を一言で言うと？',
    '権威性・実績（資格・数字・経験）は？',
    'ターゲットに提供できる情報は？',
]
for f in catch_fields:
    body_para(f'▼ {f}', bold=True, color=NAVY, size=10.5)
    input_line(lines=1)
    spacer()

body_para('▼ キャッチコピーの案を3つ書く', bold=True, color=NAVY, size=10.5)
for n in ['案1', '案2', '案3']:
    body_para(f'{n}：', size=10.5)
    input_line(lines=1)

spacer()
body_para('▼ 最終決定したキャッチコピー', bold=True, color=GOLD, size=12)
input_line(lines=1)

spacer(2)

# ポジションまとめ
p_sum = doc.add_paragraph()
p_sum.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sum.paragraph_format.space_before = Pt(16)
set_para_bg(p_sum, 0xFF, 0xF8, 0xE7)
run = p_sum.add_run('ポジションまとめ（完成形）')
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = GOLD
run.font.name = 'Yu Gothic'

t_sum = doc.add_table(rows=1, cols=1)
t_sum.style = 'Table Grid'
cell_s = t_sum.rows[0].cells[0]
set_cell_bg(cell_s, 0xFF, 0xF8, 0xE7)
csp = cell_s.paragraphs[0]
csp.paragraph_format.space_before = Pt(10)
csp.paragraph_format.space_after = Pt(10)
run_s = csp.add_run(
    '私は【　　　　　　　】に悩む【　　　　　　　】に対して、\n'
    '【　　　　　　　】という方法で【　　　　　　　】を実現するサポートをします。\n\n'
    '私だけが提供できる強みは【　　　　　　　　　　　　】です。'
)
run_s.font.size = Pt(11)
run_s.font.name = 'Yu Mincho'
run_s.font.color.rgb = NAVY

spacer(2)

# 巻末
p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_final.paragraph_format.space_before = Pt(20)
add_bottom_border(p_final, 'D4AF37', 12)

p_msg = doc.add_paragraph()
p_msg.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_msg.paragraph_format.space_before = Pt(16)
run_msg = p_msg.add_run(
    'あなたの強みは、すでにあなたの中にあります。\n'
    'このシートで言語化できたら、次はそれをInstagramで発信する番です。'
)
run_msg.font.size = Pt(11)
run_msg.font.color.rgb = NAVY
run_msg.font.name = 'Yu Mincho'
run_msg.italic = True

# 保存
out = '/Users/omorimaria/lp-preview/gift1_position_sheet.docx'
doc.save(out)
print(f'Saved: {out}')
