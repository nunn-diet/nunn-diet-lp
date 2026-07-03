from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

NAVY  = RGBColor(0x1A, 0x2B, 0x4A)
PINK  = RGBColor(0xD4, 0x53, 0x7E)
GOLD  = RGBColor(0xB8, 0x86, 0x0B)
GRAY  = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_PINK = (0xFF, 0xEA, 0xF0)
LIGHT_NAVY = (0xF0, 0xF2, 0xF8)
LIGHT_GOLD = (0xFF, 0xF8, 0xE7)
LIGHT_GREEN = (0xF0, 0xF8, 0xF0)

def set_cell_bg(cell, r, g, b):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{r:02X}{g:02X}{b:02X}')
    tcPr.append(shd)

def set_para_bg(para, r, g, b):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), f'{r:02X}{g:02X}{b:02X}')
    pPr.append(shd)

def add_bottom_border(para, color_hex='DDDDDD', size=6):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(size))
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.space_before = Pt(1)

def category_header(icon, title, subtitle, bg_rgb, text_color=WHITE):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, *bg_rgb)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after = Pt(8)
    run1 = cp.add_run(f'{icon}  {title}\n')
    run1.bold = True
    run1.font.size = Pt(14)
    run1.font.color.rgb = text_color
    run1.font.name = 'Yu Gothic'
    run2 = cp.add_run(subtitle)
    run2.font.size = Pt(9.5)
    run2.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run2.font.name = 'Yu Mincho'
    doc.add_paragraph()

def neta_card(num, title, category_bg, hint, example):
    """ネタカード"""
    # 番号+タイトル行
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(8)
    p_title.paragraph_format.space_after = Pt(2)
    run_num = p_title.add_run(f'No.{num:02d}  ')
    run_num.bold = True
    run_num.font.size = Pt(10)
    run_num.font.color.rgb = RGBColor(*category_bg)
    run_num.font.name = 'Yu Gothic'
    run_t = p_title.add_run(title)
    run_t.bold = True
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = NAVY
    run_t.font.name = 'Yu Gothic'

    # ヒント
    p_hint = doc.add_paragraph()
    p_hint.paragraph_format.left_indent = Cm(0.5)
    p_hint.paragraph_format.space_after = Pt(2)
    run_h_label = p_hint.add_run('💡 ネタのヒント：')
    run_h_label.bold = True
    run_h_label.font.size = Pt(9)
    run_h_label.font.color.rgb = GRAY
    run_h_label.font.name = 'Yu Gothic'
    run_h = p_hint.add_run(hint)
    run_h.font.size = Pt(9.5)
    run_h.font.color.rgb = GRAY
    run_h.font.name = 'Yu Mincho'

    # 投稿例
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    cell = t.rows[0].cells[0]
    set_cell_bg(cell, 0xFA, 0xFA, 0xFA)
    cp = cell.paragraphs[0]
    cp.paragraph_format.left_indent = Cm(0.3)
    cp.paragraph_format.space_before = Pt(4)
    cp.paragraph_format.space_after = Pt(4)
    run_label = cp.add_run('📝 投稿例：')
    run_label.bold = True
    run_label.font.size = Pt(9)
    run_label.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run_label.font.name = 'Yu Gothic'
    run_ex = cp.add_run(f'\n{example}')
    run_ex.font.size = Pt(9.5)
    run_ex.font.color.rgb = NAVY
    run_ex.font.name = 'Yu Mincho'
    run_ex.italic = True

    add_bottom_border(doc.add_paragraph(), 'EEEEEE', 4)

# ===================== 表紙画像 =====================
COVER_IMAGE = '/Volumes/SSD-PHPU3A/クラウドコード真里亜/ビジネスLP用/ビジネスLP/ChatGPT Image 2026年6月18日 15_15_10.png'
p_cover = doc.add_paragraph()
p_cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cover.paragraph_format.space_before = Pt(0)
p_cover.paragraph_format.space_after = Pt(0)
run_cover = p_cover.add_run()
run_cover.add_picture(COVER_IMAGE, width=Cm(16))
doc.add_page_break()

# ===================== 表紙テキスト =====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(30)
run = p.add_run('フォロワー0から')
run.font.size = Pt(13)
run.font.color.rgb = PINK
run.bold = True
run.font.name = 'Yu Gothic'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run('「この人に習いたい」と言われるまでに')
run2.font.size = Pt(16)
run2.bold = True
run2.font.color.rgb = NAVY
run2.font.name = 'Yu Gothic'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
run3 = p3.add_run('私がやった')
run3.font.size = Pt(15)
run3.bold = True
run3.font.color.rgb = NAVY
run3.font.name = 'Yu Gothic'

p4 = doc.add_paragraph()
p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
run4 = p4.add_run('自己開示ネタ 30選')
run4.font.size = Pt(26)
run4.bold = True
run4.font.color.rgb = GOLD
run4.font.name = 'Yu Gothic'

spacer()

p_line = doc.add_paragraph()
add_bottom_border(p_line, 'D4537E', 10)

spacer()

# 3要素バッジ
t_badge = doc.add_table(rows=1, cols=3)
t_badge.style = 'Table Grid'
badge_data = [
    ('共感されるストーリーの作り方', (0xFF, 0xEA, 0xF0), PINK),
    ('信頼につながる伝え方のコツ', (0xF0, 0xF2, 0xF8), NAVY),
    ('あなたらしく選ばれ続ける発信へ', (0xFF, 0xF8, 0xE7), RGBColor(0xB8, 0x86, 0x0B)),
]
for i, (text, bg, tc) in enumerate(badge_data):
    cell = t_badge.rows[0].cells[i]
    set_cell_bg(cell, *bg)
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(8)
    cp.paragraph_format.space_after = Pt(8)
    run = cp.add_run(text)
    run.font.size = Pt(9)
    run.font.color.rgb = tc
    run.font.name = 'Yu Mincho'
    run.bold = True

spacer(2)

p_point = doc.add_paragraph()
p_point.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_point = p_point.add_run(
    '「私だから話したい」と思ってもらえる\n心の距離をぐっと縮めるリアルなネタ集'
)
run_point.font.size = Pt(11)
run_point.font.color.rgb = NAVY
run_point.font.name = 'Yu Mincho'
run_point.bold = True

spacer(2)

p_intro = doc.add_paragraph()
run_intro = p_intro.add_run(
    'この資料は、Instagram・ストーリーズで使える「自己開示ネタ」を30個まとめたものです。\n'
    '発信者として「信頼」「共感」「ファン化」を生み出すためには、ノウハウや知識だけでなく、\n'
    'あなた自身の人間らしい部分を開示することが不可欠です。\n\n'
    'ネタはファン化の3要素【憧れ・共感・応援】に対応する3カテゴリーに分類しています。\n'
    '自分のアカウントに足りない要素のネタを積極的に取り入れてみてください。'
)
run_intro.font.size = Pt(10)
run_intro.font.color.rgb = GRAY
run_intro.font.name = 'Yu Mincho'

# 改ページ
doc.add_page_break()

# ===================== CATEGORY 1: 共感（ピンク）=====================
category_header('♡', 'CATEGORY 1　共感ネタ10選', '「この人と同じだ」「わかる！」と感じてもらえるネタ', (0xD4, 0x53, 0x7E))

kyokan_neta = [
    (
        'ダイエット失敗の黒歴史を語る',
        'これまで何回ダイエットに失敗したか・どんな方法で失敗したかをリアルに話す',
        '「実は私、過去に〇〇ダイエットで3kg増えたことがあって…今だから言える失敗談です笑」'
    ),
    (
        '食欲に負けた日の話',
        'ダイエット中なのに食べすぎてしまった正直な体験談',
        '「昨日、子供のチョコを3枚食べてしまいました。もう仕方ない、明日から切り替える！」'
    ),
    (
        '自分に自信がなかった頃の話',
        '体型・外見・性格など、自分を好きになれなかった時期のリアルなエピソード',
        '「鏡を見るのが怖かった時期がありました。今でも思い出すとちょっと胸が痛くなります」'
    ),
    (
        '育児と自分のやりたいことの葛藤',
        '子育て中に自分を後回しにしてしまう罪悪感や、やりたいことができないもどかしさ',
        '「子供が寝た後だけが自分の時間。でも疲れて何もできなかった昨日…そんな日ありませんか？」'
    ),
    (
        '夫婦関係のリアルな悩み',
        '夫とのコミュニケーションがうまくいかなかった時期・すれ違いのエピソード',
        '「産後、夫と全然話さなくなって。それが体型より先に崩れたものだったかもしれない」'
    ),
    (
        '「やらなきゃ」と思いつつ動けない話',
        '頭でわかっているのに行動できないもどかしさ・先延ばし癖',
        '「毎朝「今日こそ運動する」って思うのに、夜には「まぁいいか」ってなってた。あの頃の私に言いたい」'
    ),
    (
        '比べてしまって落ち込んだ話',
        '他の人と自分を比べてしまい、自己嫌悪になった経験',
        '「SNSで同い年の人のビフォーアフターを見てため息をついていた日があります。今思えば…」'
    ),
    (
        '朝が弱くてルーティンが続かない話',
        '健康的な朝習慣を作ろうとしたけど続かなかった失敗談',
        '「早起きして白湯を飲む生活、3日で終わりました。でも今は〇〇に変えたら続いてる！」'
    ),
    (
        '自分を責めてしまう癖の話',
        'うまくいかないとすぐ自己否定してしまうパターン・それを変えた気づき',
        '「ちょっとでも食べすぎると「また失敗した」って思っていた。それ自体がダイエットを失敗させていた」'
    ),
    (
        '「もう歳だから」と諦めかけた話',
        '年齢を言い訳にしていた自分・それを変えたきっかけ',
        '「30代後半になって「もう若くないし」って半分諦めていました。でも変われた。それを証明したかった」'
    ),
]

for i, (title, hint, example) in enumerate(kyokan_neta, 1):
    neta_card(i, title, (0xD4, 0x53, 0x7E), hint, example)

doc.add_page_break()

# ===================== CATEGORY 2: 憧れ（ネイビー）=====================
category_header('★', 'CATEGORY 2　憧れネタ10選', '「この人すごい」「こうなりたい」と感じてもらえるネタ', (0x1A, 0x2B, 0x4A))

akogare_neta = [
    (
        '実績・ビフォーアフターを数字で見せる',
        '体重・ウエスト・期間など具体的な数字で変化を示す',
        '「1ヶ月半でマイナス5kg、ウエストマイナス7cm。これを達成できた3つの理由を話します」'
    ),
    (
        '資格・専門知識をさりげなく出す',
        '取得した資格・学んだこと・専門家としての視点を自然に伝える',
        '「ファスティング指導士として学んでわかった、空腹感のウソについて話します」'
    ),
    (
        'クライアントの成功事例を紹介する',
        '許可を取った上でモニターさんの変化を具体的に紹介する（数字・エピソード）',
        '「先月サポートした〇〇さん、3週間でウエスト4cm減！彼女がやったことはたった1つだけでした」'
    ),
    (
        '体型だけでなく、人生が変わった話',
        '痩せたことで自信がついて行動が変わった・夫婦関係が改善したなどのライフチェンジ系',
        '「体重が落ちたとき、夫が「きれいになったね」と言ってくれた。その言葉がすべてを変えました」'
    ),
    (
        '普段の食事・習慣を自然に見せる',
        '日常の食事・運動・生活習慣を羨ましいと思ってもらえる形で発信',
        '「今日の朝ごはん。これで満足度バツグン・太らない組み合わせです」'
    ),
    (
        '勉強・インプットの様子を見せる',
        '本・セミナー・学習への姿勢を見せることで向上心・専門性をアピール',
        '「今日読んでいた本。腸内環境とダイエットの関係、面白かったので近いうちにまとめます」'
    ),
    (
        '「以前の自分とは別人」な今の姿',
        '変化前・変化後を対比させて、今の自分への信頼感を高める',
        '「3年前の私が見たら信じられないと思う。自分の体が好きになれた日が来るとは思っていなかった」'
    ),
    (
        '講師として教えた内容の一部をチラ見せ',
        '講座や個別指導で話した内容・クライアントに伝えていることのダイジェスト',
        '「先日のセッションでお伝えした「リバウンドしない人の共通点」、ここだけちょっと話しますね」'
    ),
    (
        '夫婦関係が改善した具体的なエピソード',
        '体型改善によって自己肯定感が上がり、夫婦のコミュニケーションが変わったリアルな話',
        '「体型が変わり始めた頃から、夫への態度が変わっていた。自分でも気づかなかった変化でした」'
    ),
    (
        '「やって正解だった」と心から思うこと',
        '過去の選択・努力・学びが今につながっていると思えるエピソード',
        '「ファスティングを学んだとき、正直半信半疑でした。でも今は、あれが人生を変えたと思っています」'
    ),
]

for i, (title, hint, example) in enumerate(akogare_neta, 11):
    neta_card(i, title, (0x1A, 0x2B, 0x4A), hint, example)

doc.add_page_break()

# ===================== CATEGORY 3: 応援（ゴールド）=====================
category_header('▲', 'CATEGORY 3　応援ネタ10選', '「頑張れ！」「応援したい」と感じてもらえるネタ', (0x85, 0x4F, 0x0B))

ouen_neta = [
    (
        '資格の勉強を始めた報告',
        '新しいことに挑戦するドキドキ感・理由・目標を正直に話す',
        '「今日から○○の勉強を始めました。取れたらどんなことができるようになるか、楽しみで仕方ない」'
    ),
    (
        '試験前のドキドキを発信する',
        '緊張感・不安・期待をリアルタイムで共有してフォロワーを巻き込む',
        '「明日が試験日です。準備はしてきたけど、やっぱり緊張する。応援してもらえると嬉しいです」'
    ),
    (
        '合格発表を一緒に喜ぶ',
        '結果が出たときの生の感情・感謝・これからへの思いを全力で共有する',
        '「合格しました！！！一緒にドキドキしてくれていたみなさん、本当にありがとうございます涙」'
    ),
    (
        '初めてのモニター募集の裏側',
        '初クライアントを迎えるドキドキ・不安・やりがいをリアルに話す',
        '「初めてのモニターさんのサポートが始まりました。責任感と緊張で昨夜は眠れなかったです笑」'
    ),
    (
        '食欲に正直すぎた日の話',
        '完璧じゃない自分の日常をさらけ出すことで親近感・応援される存在に',
        '「すごいお腹が空いてチョコを食べたんだけど、その後ラーメンも食べたくなって…結局ラーメンも食べちゃいました笑 こういう日ありませんか？」'
    ),
    (
        '「これをやろうと決めた」宣言',
        '目標・挑戦・決意を宣言して、フォロワーに見届けてもらう',
        '「今月中にリールを8本投稿すると決めました。見ててください！」'
    ),
    (
        'モニターさんの変化を喜ぶ',
        'クライアントの小さな変化・頑張りをキラキラした目線で紹介（許可必須）',
        '「今日のセッションで〇〇さんが「鏡を見るのが楽しくなってきた」と言ってくれた。泣きそうでした」'
    ),
    (
        '悩みながら発信している正直な話',
        '「これでいいのかな」と迷いながら続けていること・それでも続けている理由',
        '「投稿するたびに「これで合ってるのかな」と思います。でも、誰かに届いていると信じて続けます」'
    ),
    (
        '「あなたのために」という思いを伝える',
        '発信の根底にある「誰を助けたいか」という気持ちを正直に話す',
        '「私がここで発信し続けているのは、昔の私みたいに一人で悩んでいる人に届けたいから」'
    ),
    (
        '成長を振り返る投稿',
        '1ヶ月前・半年前の自分と今を比べて「変わった部分」を共有する',
        '「1ヶ月前の私は、リールを作るのに2日かかっていた。今日は3時間で完成した。小さいけど確実に成長してる」'
    ),
]

for i, (title, hint, example) in enumerate(ouen_neta, 21):
    neta_card(i, title, (0x85, 0x4F, 0x0B), hint, example)

doc.add_page_break()

# ===================== 巻末 =====================
p_end = doc.add_paragraph()
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end.paragraph_format.space_before = Pt(30)
set_para_bg(p_end, 0xFF, 0xEA, 0xF0)
run_end = p_end.add_run('自己開示を続けるためのヒント')
run_end.bold = True
run_end.font.size = Pt(14)
run_end.font.color.rgb = PINK
run_end.font.name = 'Yu Gothic'

spacer()

tips = [
    ('01', '「完璧な発信」より「リアルな発信」', 'うまく見せようとするより、今日感じたことを素直に出す方がファンがつきます。'),
    ('02', '3要素のバランスを意識する', 'ストーリーズを投稿するたびに「これは共感・憧れ・応援のどれ？」と確認する習慣をつけましょう。'),
    ('03', 'ネガティブを価値に変える', '「こんなこと発信していいの？」と思う失敗談や弱さこそ、最も共感されるコンテンツになります。'),
    ('04', '「等身大の自分」が最強のネタになる', '飾らず、背伸びせず、今日の自分をそのまま発信することが最もファンの心に刺さります。'),
    ('05', 'キャラ設定を守り続ける', '「やらないことリスト」を手元に置いて、発信のたびに確認するクセをつけましょう。'),
]

for num, title, body in tips:
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    # 番号セル
    c1 = t.rows[0].cells[0]
    set_cell_bg(c1, 0xD4, 0x53, 0x7E)
    t.columns[0].width = Cm(1.5)
    cp1 = c1.paragraphs[0]
    cp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp1.paragraph_format.space_before = Pt(10)
    cp1.paragraph_format.space_after = Pt(10)
    run1 = cp1.add_run(num)
    run1.bold = True
    run1.font.size = Pt(13)
    run1.font.color.rgb = WHITE
    run1.font.name = 'Yu Gothic'
    # 内容セル
    c2 = t.rows[0].cells[1]
    set_cell_bg(c2, 0xFF, 0xEA, 0xF0)
    cp2 = c2.paragraphs[0]
    cp2.paragraph_format.space_before = Pt(6)
    cp2.paragraph_format.space_after = Pt(2)
    cp2.paragraph_format.left_indent = Cm(0.3)
    run2a = cp2.add_run(title + '\n')
    run2a.bold = True
    run2a.font.size = Pt(10.5)
    run2a.font.color.rgb = NAVY
    run2a.font.name = 'Yu Gothic'
    run2b = cp2.add_run(body)
    run2b.font.size = Pt(9.5)
    run2b.font.color.rgb = GRAY
    run2b.font.name = 'Yu Mincho'
    doc.add_paragraph()

spacer(2)

p_final = doc.add_paragraph()
p_final.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_bg(p_final, 0x1A, 0x2B, 0x4A)
p_final.paragraph_format.space_before = Pt(14)
p_final.paragraph_format.space_after = Pt(14)
run_final = p_final.add_run(
    '「自分を見せること」が、あなたを選ばれる存在にします。\n'
    '今日から1つ、このリストのネタを発信してみてください。'
)
run_final.font.size = Pt(11)
run_final.font.color.rgb = WHITE
run_final.font.name = 'Yu Mincho'
run_final.italic = True
run_final.bold = True

out = '/Users/omorimaria/lp-preview/gift2_jikokaishineta30.docx'
doc.save(out)
print(f'Saved: {out}')
